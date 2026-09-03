"""RAG 流水线的测试。

检索/入库测试运行在零依赖组合上：``memory`` 向量后端 + ``lexical`` 嵌入，
因此整个测试套件是封闭的（无需 Ollama / Milvus / 网络），同时覆盖真实
流水线：解析 -> 切分 -> Chunk 记录 -> 检索结果 -> 重排 -> 门控。

PDF/DOCX 解析测试与 Milvus 冒烟测试需要其可选依赖（PyMuPDF /
python-docx / milvus-lite），依赖不可导入时自动跳过。
"""

import io
import shutil
import tempfile
import unittest

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from unittest.mock import patch

from rag import chat, chunking, parsers, rag_engine, reranker, services, vector_store
from rag.models import Chunk, Document, KnowledgeBase, Message, Session
from rag.tokenizer import tokenize

TEST_SETTINGS = {
    "RAG_EMBEDDING_PROVIDER": "lexical",
    "RAG_EMBEDDING_DIM": 512,
    "RAG_INDEX_DIR": tempfile.mkdtemp(prefix="rag_test_"),
    "RAG_VECTOR_BACKEND": "memory",
    "RAG_CHUNK_SIZE": 200,
    "RAG_CHUNK_OVERLAP": 40,
    "RAG_SIMILARITY_THRESHOLD": 0.0,  # lexical 得分整体偏低，默认关闭门控
}

TEST_DIR = TEST_SETTINGS["RAG_INDEX_DIR"]


class TokenizerTests(TestCase):
    def test_cjk_bigrams(self):
        self.assertEqual(tokenize("机器学习"), ["机器", "器学", "学习"])

    def test_latin_words_lowercased(self):
        self.assertIn("rag", tokenize("RAG Backend"))

    def test_mixed(self):
        toks = tokenize("用BM25做检索")
        self.assertIn("bm25", toks)
        self.assertIn("检索", toks)


class ParserTests(TestCase):
    def test_plain_text(self):
        pages = parsers.extract_text(
            b"hello world", "text/plain", "a.txt"
        )
        self.assertEqual(pages, [{"page": None, "text": "hello world"}])

    def test_html_is_stripped(self):
        raw = "<html><body><p>定价规则 &amp; 说明</p></body></html>".encode("utf-8")
        pages = parsers.extract_text(raw, "text/html", "a.html")
        self.assertEqual(len(pages), 1)
        self.assertIn("定价规则 & 说明", pages[0]["text"])
        self.assertNotIn("<p>", pages[0]["text"])

    def test_clean_text_drops_control_chars(self):
        pages = parsers.extract_text(
            b"a\r\nb\x00c\x1fd", "text/plain", "a.txt"
        )
        self.assertEqual(pages, [{"page": None, "text": "a\nbcd"}])

    def test_unsupported_format(self):
        with self.assertRaises(ValueError):
            parsers.extract_text(b"\x89PNG\r\n", "image/png", "a.png")
        with self.assertRaises(ValueError):
            parsers.extract_text(b"\x00\x00", "application/octet-stream", "a.bin")

    @unittest.skipUnless(
        __import__("importlib").util.find_spec("fitz"),
        "PyMuPDF not installed",
    )
    def test_pdf_per_page(self):
        import fitz

        doc = fitz.open()
        try:
            for text in ("Alpha beta gamma delta page one.", "Epsilon zeta page two."):
                page = doc.new_page()
                page.insert_text((72, 72), text)
            raw = doc.tobytes()
        finally:
            doc.close()

        pages = parsers.extract_text(raw, "application/pdf", "a.pdf")
        self.assertEqual(len(pages), 2)
        self.assertEqual([p["page"] for p in pages], [1, 2])
        self.assertIn("page one", pages[0]["text"])

    @unittest.skipUnless(
        __import__("importlib").util.find_spec("fitz"),
        "PyMuPDF not installed",
    )
    def test_garbage_pdf_raises(self):
        with self.assertRaises(ValueError) as ctx:
            parsers.extract_text(b"%PDF-1.4 not really a pdf", "application/pdf", "a.pdf")
        self.assertIn("PDF", str(ctx.exception))

    @unittest.skipUnless(
        __import__("importlib").util.find_spec("docx"),
        "python-docx not installed",
    )
    def test_docx_paragraphs(self):
        from docx import Document as DocxDocument

        buf = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("第一段：产品定价说明。")
        doc.add_paragraph("第二段：售后服务政策。")
        doc.save(buf)

        pages = parsers.extract_text(
            buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "a.docx"
        )
        self.assertEqual(len(pages), 1)
        self.assertIsNone(pages[0]["page"])
        self.assertIn("售后服务政策", pages[0]["text"])


class ChunkingTests(TestCase):
    def test_chunk_pages_keep_page_metadata(self):
        chunks = chunking.chunk_pages(
            [{"page": 1, "text": "这是第一页的一段较长内容，用于测试切分行为。"}],
            chunk_size=12,
            chunk_overlap=2,
        )
        self.assertGreater(len(chunks), 0)
        self.assertTrue(all(c["page"] == 1 for c in chunks))
        # 重叠窗口不应凭空生成内容。
        joined = "".join(c["text"] for c in chunks)
        self.assertIn("第一页", joined)


class RerankerTests(TestCase):
    def test_lexical_boost_reorders_candidates(self):
        # A（纯向量得分最高但无词法重合）对 B（向量中等但词法重合度高）。
        # alpha=0.4 时 B 应胜出。
        candidates = [
            {"chunk_id": 1, "score": 0.9, "text": "股票价格走势与宏观经济数据分析报告。"},
            {"chunk_id": 2, "score": 0.7, "text": "机器学习在医疗影像诊断中的应用研究进展。"},
        ]
        with override_settings(RAG_RERANK_ALPHA=0.4):
            out = reranker.rerank("机器学习应用", candidates, top_k=2)
        self.assertEqual(out[0]["chunk_id"], 2)
        self.assertIn("rerank_score", out[0])

    def test_rerank_respects_top_k(self):
        candidates = [
            {"chunk_id": i, "score": 0.8 - i * 0.1, "text": f"内容片段 number {i}"}
            for i in range(5)
        ]
        out = reranker.rerank("number", candidates, top_k=3)
        self.assertEqual(len(out), 3)


@override_settings(**TEST_SETTINGS)
class RagEngineTests(TestCase):
    def setUp(self):
        rag_engine._embed_model_instance = None  # 使引擎重新加载 lexical 嵌入
        self.addCleanup(shutil.rmtree, TEST_DIR, ignore_errors=True)

    def _make_kb(self, docs):
        kb = KnowledgeBase.objects.create(name="测试库")
        for i, text in enumerate(docs):
            doc = Document.objects.create(
                kb=kb,
                title=f"d{i}",
                file_name=f"d{i}.txt",
                content_type="text/plain",
                size=len(text),
            )
            services.ingest_document(
                doc, text.encode("utf-8"), "text/plain", f"d{i}.txt"
            )
        return kb

    def test_ingest_chunks_and_index_rebuildable(self):
        kb = self._make_kb(["机器学习与自然语言处理的研究进展"])
        self.assertEqual(Document.objects.filter(kb=kb).count(), 1)
        self.assertGreaterEqual(kb.chunks.count(), 1)

    def test_retrieval_envelope_rank_and_metadata(self):
        kb = self._make_kb(
            [
                "智能问答模块调用大语言模型生成回答，支持配置模型接口与API密钥。",
                "本产品是一款面向企业的知识管理平台，支持文档上传与全文检索。",
            ]
        )
        envelope = services.retrieve(kb, "大语言模型如何生成智能问答回答", top_k=2)
        self.assertEqual(len(envelope["results"]), 2)
        self.assertFalse(envelope["refused"])
        self.assertEqual(envelope["backend"], "memory")
        first = envelope["results"][0]
        self.assertIn("rerank_score", first)
        self.assertIn("page", first)
        self.assertIn("document_title", first)
        # 目标片段必须排在通用平台片段之前。
        self.assertEqual(first["document_title"], "d0")

    def test_retrieval_empty_kb_refused(self):
        kb = KnowledgeBase.objects.create(name="空库")
        envelope = services.retrieve(kb, "任何查询")
        self.assertTrue(envelope["refused"])
        self.assertEqual(envelope["results"], [])

    @override_settings(RAG_SIMILARITY_THRESHOLD=0.99)
    def test_below_threshold_is_refused_without_results(self):
        kb = self._make_kb(["机器学习与自然语言处理的研究进展"])
        envelope = services.retrieve(kb, "机器学习研究进展")
        self.assertTrue(envelope["refused"])
        self.assertEqual(envelope["results"], [])
        self.assertGreaterEqual(envelope["best_score"], 0.0)

    def test_delete_document_removes_from_retrieval(self):
        kb = self._make_kb(["机器翻译与对话系统综述内容"])
        envelope = services.retrieve(kb, "机器翻译")
        self.assertEqual(envelope["count"], 1)
        doc = Document.objects.get(kb=kb)
        doc.delete()
        rag_engine.invalidate_kb_index(kb)
        envelope = services.retrieve(kb, "机器翻译")
        self.assertTrue(envelope["refused"])
        self.assertEqual(envelope["results"], [])

    def test_unsupported_format_marks_failed(self):
        kb = KnowledgeBase.objects.create(name="另库")
        doc = Document.objects.create(kb=kb, title="d", file_name="d.xlsx")
        with self.assertRaises(ValueError):
            parsers.extract_text(b"PK\x03\x04fake", "application/octet-stream", "d.xlsx")
        with self.assertRaises(ValueError):
            services.ingest_document(doc, b"%PDF-x", "application/pdf", "d.pdf")
        doc.refresh_from_db()
        self.assertEqual(doc.status, Document.Status.FAILED)
        self.assertIn("PDF", doc.error)

    @unittest.skipUnless(
        __import__("importlib").util.find_spec("fitz"),
        "PyMuPDF not installed",
    )
    def test_pdf_page_metadata_lands_on_chunks(self):
        import fitz

        doc = fitz.open()
        try:
            page = doc.new_page()
            page.insert_text((72, 72), "Pricing rules page one content alpha beta.")
            raw = doc.tobytes()
        finally:
            doc.close()

        kb = KnowledgeBase.objects.create(name="pdf库")
        uploaded = Document.objects.create(kb=kb, title="p.pdf", file_name="p.pdf")
        services.ingest_document(uploaded, raw, "application/pdf", "p.pdf")
        uploaded.refresh_from_db()
        self.assertEqual(uploaded.status, Document.Status.READY)
        chunk_rows = list(Chunk.objects.filter(document=uploaded))
        self.assertGreater(len(chunk_rows), 0)
        self.assertTrue(all(c.page == 1 for c in chunk_rows))
        self.assertIn("第 1 页", uploaded.text)


@override_settings(**TEST_SETTINGS)
class ApiTests(TestCase):
    def setUp(self):
        rag_engine._embed_model_instance = None
        self.addCleanup(shutil.rmtree, TEST_DIR, ignore_errors=True)

    def _kb(self):
        return KnowledgeBase.objects.create(name="接口库", description="api")

    def _upload(self, kb, name="doc.txt", content="智能问答系统介绍文本。"):
        url = reverse("rag:document-list")
        resp = self.client.post(
            url,
            {
                "file": SimpleUploadedFile(name, content.encode("utf-8")),
                "kb": kb.id,
            },
        )
        return resp

    def test_kb_crud_and_search_endpoint(self):
        url = reverse("rag:knowledge-base-list")
        resp = self.client.post(url, {"name": "新库", "description": "x"}, content_type="application/json")
        self.assertEqual(resp.status_code, 201)
        kb = KnowledgeBase.objects.get(name="新库")

        self._upload(kb)
        search_url = reverse("rag:knowledge-base-search", kwargs={"pk": kb.id})
        resp = self.client.post(
            search_url, {"query": "智能问答", "top_k": 5}, content_type="application/json"
        )
        self.assertEqual(resp.status_code, 200)
        payload = resp.json()
        self.assertFalse(payload["refused"])
        self.assertEqual(payload["count"], len(payload["results"]))
        self.assertIn("rerank_score", payload["results"][0])
        self.assertIn("threshold", payload)
        self.assertIn("best_score", payload)

    def test_documents_filtered_by_kb_and_delete(self):
        kb1 = self._kb()
        kb2 = KnowledgeBase.objects.create(name="另一个库")
        self._upload(kb1)
        self._upload(kb2)

        resp = self.client.get(reverse("rag:document-list"), {"kb": kb1.id})
        self.assertEqual(len(resp.json()), 1)
        doc = Document.objects.get(kb=kb1)
        resp = self.client.delete(reverse("rag:document-detail", kwargs={"pk": doc.id}))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(Document.objects.filter(kb=kb1).count(), 0)

    def test_upload_missing_kb_rejected(self):
        resp = self.client.post(
            reverse("rag:document-list"),
            {"file": SimpleUploadedFile("a.txt", b"hello")},
        )
        self.assertEqual(resp.status_code, 400)


@override_settings(**TEST_SETTINGS)
class ChatSseTests(TestCase):
    def setUp(self):
        rag_engine._embed_model_instance = None
        self.addCleanup(shutil.rmtree, TEST_DIR, ignore_errors=True)

    def _kb_with_doc(self, text):
        kb = KnowledgeBase.objects.create(name="对话库")
        doc = Document.objects.create(kb=kb, title="d", file_name="d.txt")
        services.ingest_document(doc, text.encode("utf-8"), "text/plain", "d.txt")
        return kb

    def _chat(self, kb, query, session_id=None):
        url = reverse("rag:knowledge-base-chat", kwargs={"pk": kb.id})
        payload = {"query": query}
        if session_id:
            payload["session_id"] = session_id
        return self.client.post(url, payload, content_type="application/json")

    def test_streamed_answer_persists_messages(self):
        kb = self._kb_with_doc("产品定价规则说明，包含会员折扣与批量优惠。")
        # 注意：StreamingHttpResponse 惰性求值，生成器在迭代响应时才执行，
        # 因此 mock 与断言必须保持在同一个 with 作用域内。
        with patch(
            "rag.llm.generate_answer_stream",
            return_value=iter(["会员折扣为八折。", "批量采购另有优惠。"]),
        ):
            resp = self._chat(kb, "会员折扣是多少")
            self.assertEqual(resp.status_code, 200)
            body = b"".join(resp.streaming_content).decode("utf-8")
        self.assertIn('event: sources', body)
        self.assertIn('event: chunk', body)
        self.assertIn('会员折扣为八折。', body)
        self.assertIn('event: done', body)

        session = Session.objects.get(kb=kb)
        messages = list(session.messages.order_by("id"))
        self.assertEqual(len(messages), 2)
        self.assertEqual(messages[0].role, Message.Role.USER)
        self.assertEqual(messages[1].role, Message.Role.ASSISTANT)
        self.assertGreater(len(messages[1].sources), 0)
        self.assertIn("会员折扣为八折。", messages[1].content)

    def test_refused_query_skips_llm(self):
        kb = self._kb_with_doc("机器学习与自然语言处理研究综述。")
        with override_settings(RAG_SIMILARITY_THRESHOLD=0.99):
            with patch("rag.llm.generate_answer_stream") as stream_mock:
                resp = self._chat(kb, "机器学习")
                # 惰性流：override_settings 与 mock 必须覆盖到迭代完成。
                body = b"".join(resp.streaming_content).decode("utf-8")
        self.assertIn("event: refused", body)
        self.assertIn("event: done", body)
        stream_mock.assert_not_called()

        session = Session.objects.get(kb=kb)
        assistant = session.messages.get(role=Message.Role.ASSISTANT)
        self.assertEqual(assistant.content, chat.REFUSED_ANSWER)
        self.assertEqual(assistant.sources, [])

    def test_session_id_must_belong_to_kb(self):
        kb = self._kb_with_doc("内容")
        other = KnowledgeBase.objects.create(name="别的库")
        session = Session.objects.create(kb=other, title="x")
        resp = self._chat(kb, "问题", session_id=session.id)
        self.assertEqual(resp.status_code, 404)

    def test_sessions_endpoints(self):
        kb = self._kb_with_doc("内容")
        url = reverse("rag:session-list")
        resp = self.client.post(url, {"kb": kb.id, "title": "会话一"}, content_type="application/json")
        self.assertEqual(resp.status_code, 201)
        session_id = resp.json()["id"]

        resp = self.client.get(url, {"kb": kb.id})
        self.assertEqual(len(resp.json()), 1)

        Message.objects.create(session_id=session_id, role="user", content="hi")
        resp = self.client.get(
            reverse("rag:session-messages", kwargs={"pk": session_id})
        )
        payload = resp.json()
        self.assertEqual(payload["count"], 1)
        self.assertEqual(payload["messages"][0]["content"], "hi")

        resp = self.client.delete(reverse("rag:session-detail", kwargs={"pk": session_id}))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(Session.objects.count(), 0)


@unittest.skipUnless(
    __import__("importlib").util.find_spec("milvus_lite"),
    "milvus-lite not installed",
)
@override_settings(**TEST_SETTINGS)
class MilvusSmokeTests(TestCase):
    """Milvus Lite 真实往返测试：创建 collection -> 插入 -> 搜索 ->
    条件删除 -> 删除。缺少 milvus-lite 时自动跳过。"""

    def test_roundtrip(self):
        tmp = tempfile.mkdtemp(prefix="rag_milvus_")
        self.addCleanup(shutil.rmtree, tmp, ignore_errors=True)
        backend = vector_store.MilvusBackend(uri=str(tmp) + "/smoke.db")
        try:
            backend.ping()
        except RuntimeError as exc:  # pragma: no cover - 与环境相关
            self.skipTest(f"milvus-lite cannot start: {exc}")

        from rag.rag_engine import LexicalEmbedding

        embed = LexicalEmbedding(dim=512)
        kb = KnowledgeBase.objects.create(name="milvus-smoke")
        doc = Document.objects.create(kb=kb, title="m", file_name="m.txt")

        rows = [
            {"chunk_id": 101, "document_id": doc.id, "page": None, "text": "产品定价规则", "embedding": embed._embed("产品定价规则")},
            {"chunk_id": 102, "document_id": doc.id, "page": 3, "text": "会员折扣方案", "embedding": embed._embed("会员折扣方案")},
        ]
        backend.upsert(kb, rows)
        hits = backend.search(kb, None, embed._embed("会员折扣"), top_k=5)
        ids = {h["chunk_id"] for h in hits}
        self.assertIn(102, ids)
        self.assertGreater(hits[0]["score"], 0.0)

        # 模拟引擎重启：同一数据目录重新打开后，collection 处于 released
        # 状态，search 必须经过 load() 才能命中（回归 vector_store 曾缺失
        # load 导致 "call load() before search" 的问题）。
        backend.close()
        backend2 = vector_store.MilvusBackend(uri=backend._uri)
        hits = backend2.search(kb, None, embed._embed("会员折扣"), top_k=5)
        ids = {h["chunk_id"] for h in hits}
        self.assertIn(102, ids)

        backend2.delete_document(kb, doc.id)
        hits = backend2.search(kb, None, embed._embed("会员折扣"), top_k=5)
        self.assertEqual(hits, [])
        backend2.clear_kb(kb.id)
